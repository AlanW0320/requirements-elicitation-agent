import os
import re
import json
import sqlite3
import pandas as pd
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()
client     = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
AGENT_MODEL = "gpt-4o"


# ══════════════════════════════════════════════════════════════════════════
# REGISTRY LOADER
# ══════════════════════════════════════════════════════════════════════════

def load_requirements_from_registry(
    db_path      = './data/requirements_registry.db',
    project_name = None
):
    """
    Load all classified requirements from the SQLite registry.
    Excludes Ambiguous entries — everything else is guaranteed classified.
    Returns a DataFrame and a dict grouped by label.
    """
    conn = sqlite3.connect(db_path)

    if project_name:
        rows = conn.execute('''
            SELECT final_text, label, confidence, status
            FROM requirements
            WHERE project_name = ?
            AND label NOT IN ("Ambiguous")
        ''', (project_name,)).fetchall()
    else:
        rows = conn.execute('''
            SELECT final_text, label, confidence, status
            FROM requirements
            WHERE label NOT IN ("Ambiguous")
        ''').fetchall()

    conn.close()

    df = pd.DataFrame(rows, columns=['text', 'label', 'confidence', 'status'])

    # Deduplicate by text
    before = len(df)
    df     = df.drop_duplicates(subset=['text']).reset_index(drop=True)
    after  = len(df)
    if before != after:
        print(f"  Removed {before - after} duplicate requirements")

    # Group by label
    grouped = {}
    for label in df['label'].unique():
        grouped[label] = df[df['label'] == label]['text'].tolist()

    print(f"Total requirements loaded: {len(df)}")
    print(f"\nBy category:")
    for label, reqs in grouped.items():
        print(f"  {label:<25}: {len(reqs)}")

    return df, grouped


def load_flagged_requirements(
    db_path      = './data/requirements_registry.db',
    project_name = None
):
    """Load requirements that were flagged as Ambiguous — for appendix use."""
    conn = sqlite3.connect(db_path)
    if project_name:
        rows = conn.execute('''
            SELECT final_text FROM requirements
            WHERE project_name = ?
            AND label IN ("Ambiguous", "Requires Manual Review")
        ''', (project_name,)).fetchall()
    else:
        rows = conn.execute('''
            SELECT final_text FROM requirements
            WHERE label IN ("Ambiguous", "Requires Manual Review")
        ''').fetchall()
    conn.close()
    return [row[0] for row in rows]


# ══════════════════════════════════════════════════════════════════════════
# HELPER UTILITIES
# ══════════════════════════════════════════════════════════════════════════

def strip_gpt_section_header(text, section_number):
    """
    Remove the first line of GPT output if it repeats the section heading.
    e.g. removes '2. Overall Description' or '1. Introduction'.
    """
    lines      = text.strip().split('\n')
    if not lines:
        return text
    first_line = lines[0].strip()
    if re.match(rf'^{re.escape(str(section_number))}[\.\s]', first_line):
        lines = lines[1:]
    while lines and not lines[0].strip():
        lines.pop(0)
    return '\n'.join(lines).strip()


def is_requirement_id(line):
    """
    Detect requirement ID lines like FR-001, NFR_Security-001, NFRSecurity-001.
    Strips leading bullets before checking.
    """
    cleaned = line.strip().lstrip('·•-\u2013 \t')
    return bool(re.match(r'^(FR|NFR[_]?[A-Za-z]*)-\d{3}$', cleaned.strip()))


def clean_id_line(line):
    """Strip leading bullet characters from ID lines."""
    return line.strip().lstrip('·•-\u2013 \t').strip()


# ══════════════════════════════════════════════════════════════════════════
# SECTION GENERATORS
# ══════════════════════════════════════════════════════════════════════════

def generate_introduction_section(context, grouped):
    """Generate Section 1 — Introduction."""
    total_reqs = sum(len(v) for v in grouped.values())
    fr_count   = len(grouped.get('FR', []))
    nfr_count  = total_reqs - fr_count

    prompt = f"""You are a software requirements engineer writing a formal SRS document.

Write Section 1 (Introduction) of an IEEE 830 Software Requirements Specification.

Project details:
- Name        : {context['project_name']}
- Description : {context['description']}
- Scope       : {context['scope']}
- Intended users: {', '.join(context['intended_users'])}
- Total requirements: {total_reqs} ({fr_count} functional, {nfr_count} non-functional)

Write these subsections in clear professional technical writing:
1.1 Purpose
1.2 Scope
1.3 Definitions and Abbreviations
1.4 Intended Audience
1.5 Document Overview

IMPORTANT: Do NOT include "1. Introduction" as a heading — the script adds headings automatically.
Use formal IEEE 830 language. Plain paragraphs only — no markdown, no bullet points."""

    response = client.chat.completions.create(
        model      = AGENT_MODEL,
        messages   = [
            {"role": "system", "content": "You are a technical writer producing IEEE 830 SRS documents."},
            {"role": "user",   "content": prompt}
        ],
        max_tokens  = 1000,
        temperature = 0.2
    )
    return response.choices[0].message.content.strip()


def generate_overall_description(context):
    """Generate Section 2 — Overall Description."""
    prompt = f"""Write Section 2 (Overall Description) of an IEEE 830 SRS document.

Project: {context['project_name']}
Description: {context['description']}

Include these subsections:
2.1 Product Perspective
2.2 Product Functions (high-level summary only)
2.3 User Characteristics
2.4 Constraints
2.5 Assumptions and Dependencies

IMPORTANT: Do NOT include "2. Overall Description" as a heading — start directly
with the introductory paragraph followed by subsections.
Use formal IEEE 830 language. Plain paragraphs only — no markdown."""

    response = client.chat.completions.create(
        model      = AGENT_MODEL,
        messages   = [
            {"role": "system", "content": "You are a technical writer producing IEEE 830 SRS documents."},
            {"role": "user",   "content": prompt}
        ],
        max_tokens  = 800,
        temperature = 0.2
    )
    return response.choices[0].message.content.strip()


def generate_functional_requirements(fr_list):
    """Generate Section 3.1 — Functional Requirements with unique IDs."""
    if not fr_list:
        return "No functional requirements have been classified."

    # Pre-assign IDs so GPT cannot add or split requirements
    reqs_text = '\n'.join([
        f"FR-{str(i+1).zfill(3)}: {req}"
        for i, req in enumerate(fr_list)
    ])

    prompt = f"""You are writing Section 3.1 of an IEEE 830 SRS.

You have EXACTLY {len(fr_list)} functional requirements to document.
Do NOT add, split, merge, or invent any requirements.
Document ONLY the {len(fr_list)} requirements listed below.

Requirements:
{reqs_text}

For each requirement keep its assigned ID and add:
- Description: copy the requirement text exactly as given
- Priority: High / Medium / Low
- Acceptance Criterion: one specific testable criterion

CRITICAL FORMATTING RULES:
1. The ID line (FR-001) must be on its own line with NO bullet, NO dash, NO dot before it
2. Do NOT put a bullet point or dash before any ID line
3. Separate each requirement with exactly one blank line
4. No markdown

Format EXACTLY:
FR-001
Description: [exact requirement text]
Priority: High
Acceptance Criterion: [one testable criterion]

IMPORTANT: Do NOT include a section heading — start directly with FR-001."""

    response = client.chat.completions.create(
        model      = AGENT_MODEL,
        messages   = [
            {
                "role"   : "system",
                "content": (
                    "You are a requirements engineer. Document ONLY the "
                    "requirements provided. Never add new ones. "
                    "Never put bullets before requirement IDs."
                )
            },
            {"role": "user", "content": prompt}
        ],
        max_tokens  = 2000,
        temperature = 0
    )
    return response.choices[0].message.content.strip()


def generate_nfr_section(grouped):
    """Generate Section 3.2 — Non-Functional Requirements grouped by category."""
    nfr_categories = {
        k: v for k, v in grouped.items()
        if k.startswith('NFR_') and v
    }

    if not nfr_categories:
        return "No non-functional requirements have been classified."

    category_names = {
        'NFR_Performance'     : '3.2.1 Performance Requirements',
        'NFR_Security'        : '3.2.2 Security Requirements',
        'NFR_Usability'       : '3.2.3 Usability Requirements',
        'NFR_Reliability'     : '3.2.4 Reliability Requirements',
        'NFR_Maintainability' : '3.2.5 Maintainability Requirements',
        'NFR_Scalability'     : '3.2.6 Scalability Requirements',
        'NFR_Portability'     : '3.2.7 Portability Requirements',
        'NFR_Operational'     : '3.2.8 Operational Requirements',
        'NFR_Legal'           : '3.2.9 Legal and Compliance Requirements',
        'NFR_LookAndFeel'     : '3.2.10 Look and Feel Requirements',
        'NFR_Other'           : '3.2.11 Other Non-Functional Requirements',
    }

    all_nfr_text = ""

    for category, reqs in nfr_categories.items():
        if not reqs:
            continue

        section_name = category_names.get(category, category)
        prefix       = category.replace('NFR_', 'NFR')

        # Pre-assign IDs
        reqs_text = '\n'.join([
            f"{prefix}-{str(i+1).zfill(3)}: {req}"
            for i, req in enumerate(reqs)
        ])

        prompt = f"""Write the {section_name} subsection of an IEEE 830 SRS.

You have EXACTLY {len(reqs)} requirements to document.
Do NOT add, split, merge, or invent any requirements.
Document ONLY these {len(reqs)} requirements.

Requirements:
{reqs_text}

For each keep its assigned ID and add:
- Description: copy the requirement text exactly
- Target: measurable criterion (infer reasonable values if missing)
- Verification: how to verify this requirement

CRITICAL FORMATTING RULES:
1. The ID line ({prefix}-001) must be on its own line with NO bullet, NO dash before it
2. Do NOT put a bullet point or dash before any ID line
3. Separate each requirement with exactly one blank line
4. No markdown

Format EXACTLY:
{prefix}-001
Description: [exact text]
Target: [measurable criterion]
Verification: [verification method]

IMPORTANT: Do NOT include the section heading "{section_name}" in your response.
Start directly with the first requirement ID."""

        response = client.chat.completions.create(
            model      = AGENT_MODEL,
            messages   = [
                {
                    "role"   : "system",
                    "content": (
                        "You are a requirements engineer. Document ONLY the "
                        "requirements provided. Never add new ones. "
                        "Never put bullets before requirement IDs."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            max_tokens  = 1200,
            temperature = 0
        )

        all_nfr_text += f"\n{section_name}\n\n"
        all_nfr_text += response.choices[0].message.content.strip()
        all_nfr_text += "\n\n"

    return all_nfr_text


# ══════════════════════════════════════════════════════════════════════════
# DOCX BUILDER
# ══════════════════════════════════════════════════════════════════════════

def build_srs_docx(
    context,
    section1,
    section2,
    section3_fr,
    section3_nfr,
    fr_list,
    flagged = None
):
    """
    Build a professionally formatted IEEE 830 SRS .docx file.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Helper functions ──────────────────────────────────────────────────
    def add_heading(text, level=1):
        heading           = doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if heading.runs:
            heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        return heading

    def add_paragraph(text, bold=False, size=11):
        para = doc.add_paragraph()
        run  = para.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        return para

    def add_req_id(text):
        """Add a requirement ID line — blue bold, no bullet."""
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold      = True
        run.font.size      = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        return p

    def add_req_detail(text):
        """Add a requirement detail line — bulleted."""
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25)
        run  = para.add_run(text.lstrip('·•-\u2013 \t'))
        run.font.size = Pt(10.5)
        return para

    def render_requirements_block(content):
        """
        Parse and render a block of GPT-generated requirements
        into properly formatted docx elements.
        """
        for line in content.split('\n'):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            if is_requirement_id(stripped):
                add_req_id(clean_id_line(stripped))
            else:
                add_req_detail(stripped)

    # ── Title Page ────────────────────────────────────────────────────────
    title           = doc.add_heading('', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run             = title.add_run('Software Requirements Specification')
    run.font.size      = Pt(24)
    run.font.bold      = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    subtitle           = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run                = subtitle.add_run(context['project_name'])
    run.font.size      = Pt(16)
    run.font.bold      = True

    doc.add_paragraph()

    # Metadata table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.style = 'Table Grid'
    meta_data = [
        ('Version',      context['version']),
        ('Date',         context['date']),
        ('Status',       context['document_status']),
        ('Organization', context['organization']),
        ('Authors',      ', '.join(context['authors'])),
    ]
    for i, (key, value) in enumerate(meta_data):
        meta_table.rows[i].cells[0].text = key
        meta_table.rows[i].cells[1].text = value
        for para in meta_table.rows[i].cells[0].paragraphs:
            for run in para.runs:
                run.font.bold = True

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────────────────
    add_heading('Table of Contents', level=1)
    toc_items = [
        '1.0 Introduction',
        '    1.1 Purpose',
        '    1.2 Scope',
        '    1.3 Definitions and Abbreviations',
        '    1.4 Intended Audience',
        '    1.5 Document Overview',
        '2.0 Overall Description',
        '    2.1 Product Perspective',
        '    2.2 Product Functions',
        '    2.3 User Characteristics',
        '    2.4 Constraints',
        '    2.5 Assumptions and Dependencies',
        '3.0 Specific Requirements',
        '    3.1 Functional Requirements',
        '    3.2 Non-Functional Requirements',
    ]
    for item in toc_items:
        add_paragraph(item)

    doc.add_page_break()

    # ── Section 1: Introduction ───────────────────────────────────────────
    add_heading('1.0 Introduction', level=1)
    for line in section1.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r'^1\.\d', stripped):
            add_heading(stripped, level=2)
        else:
            add_paragraph(stripped)

    doc.add_page_break()

    # ── Section 2: Overall Description ───────────────────────────────────
    add_heading('2.0 Overall Description', level=1)
    for line in section2.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r'^2\.\d', stripped):
            add_heading(stripped, level=2)
        else:
            add_paragraph(stripped)

    doc.add_page_break()

    # ── Section 3: Specific Requirements ─────────────────────────────────
    add_heading('3.0 Specific Requirements', level=1)

    # 3.1 Functional Requirements
    add_heading('3.1 Functional Requirements', level=2)
    add_paragraph(
        f"This section defines {len(fr_list)} functional requirement(s) "
        f"elicited and classified by the AI Requirements Elicitation Agent."
    )
    doc.add_paragraph()
    render_requirements_block(section3_fr)

    doc.add_page_break()

    # 3.2 Non-Functional Requirements
    add_heading('3.2 Non-Functional Requirements', level=2)

    for line in section3_nfr.split('\n'):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph()
            continue

        # Subsection heading e.g. "3.2.1 Performance Requirements"
        if re.match(r'^3\.2\.\d+', stripped):
            add_heading(stripped, level=3)

        # Requirement ID line
        elif is_requirement_id(stripped):
            add_req_id(clean_id_line(stripped))

        # Detail line
        else:
            add_req_detail(stripped)

    # ── Appendix: Flagged Requirements ───────────────────────────────────
    if flagged:
        doc.add_page_break()
        add_heading('Appendix A — Requirements Pending Manual Review', level=1)
        add_paragraph(
            f"The following {len(flagged)} requirement(s) could not be "
            f"automatically classified with sufficient confidence. "
            f"A requirements engineer should review these before finalizing the SRS."
        )
        doc.add_paragraph()
        for i, req in enumerate(flagged, 1):
            para = doc.add_paragraph(style='List Number')
            para.paragraph_format.left_indent = Inches(0.25)
            run  = para.add_run(req)
            run.font.size = Pt(10.5)

    # ── Footer note ───────────────────────────────────────────────────────
    doc.add_page_break()
    footer_para           = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run                   = footer_para.add_run(
        f"Generated by AI Requirements Elicitation Agent  —  {context['date']}"
    )
    run.font.size      = Pt(9)
    run.font.italic    = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    return doc


# ══════════════════════════════════════════════════════════════════════════
# FULL PIPELINE — PUBLIC ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════

def generate_srs_from_registry(
    project_name,
    db_path  = './data/requirements_registry.db',
    context  = None
):
    """
    Full pipeline — loads requirements from registry and generates
    a complete IEEE 830 SRS .docx file.

    Parameters:
        project_name : str   — name of the project in the registry
        db_path      : str   — path to the SQLite registry database
        context      : dict  — project metadata (name, version, description etc.)
                               If None, a default context is generated.

    Returns:
        output_path : str — path to the generated .docx file, or None on failure
    """
    print(f"\nGenerating SRS for: {project_name}")
    print("=" * 50)

    # ── Load requirements ─────────────────────────────────────────────────
    df, grouped = load_requirements_from_registry(db_path, project_name)

    if df.empty:
        print("⚠️  No classified requirements found in registry.")
        return None

    fr_list = grouped.get('FR', [])
    flagged = load_flagged_requirements(db_path, project_name)

    # ── Default context if not provided ──────────────────────────────────
    if context is None:
        context = {
            "project_name"    : project_name,
            "version"         : "1.0",
            "date"            : datetime.now().strftime("%B %d, %Y"),
            "organization"    : "Software Engineering Department",
            "authors"         : ["Requirements Engineering Team"],
            "document_status" : "Draft",
            "description"     : f"Software requirements for {project_name}.",
            "intended_users"  : [
                "Software developers",
                "Requirements engineers",
                "Project managers"
            ],
            "scope"           : (
                f"This document defines the functional and non-functional "
                f"requirements for {project_name}."
            )
        }

    # ── Generate sections ─────────────────────────────────────────────────
    print("\nGenerating sections...")

    print("  [1/4] Introduction...")
    section1 = generate_introduction_section(context, grouped)
    section1 = strip_gpt_section_header(section1, 1)
    print("     Done")

    print("  [2/4] Overall Description...")
    section2 = generate_overall_description(context)
    section2 = strip_gpt_section_header(section2, 2)
    print("     Done")

    print(f"  [3/4] Functional Requirements ({len(fr_list)} FRs)...")
    section3_fr = generate_functional_requirements(fr_list)
    section3_fr = strip_gpt_section_header(section3_fr, 3)
    print("     Done")

    print("  [4/4] Non-Functional Requirements...")
    section3_nfr = generate_nfr_section(grouped)
    print("     Done")

    # ── Build docx ────────────────────────────────────────────────────────
    print("\nBuilding document...")
    doc = build_srs_docx(
        context      = context,
        section1     = section1,
        section2     = section2,
        section3_fr  = section3_fr,
        section3_nfr = section3_nfr,
        fr_list      = fr_list,
        flagged      = flagged if flagged else None
    )

    # ── Save ──────────────────────────────────────────────────────────────
    os.makedirs('./outputs', exist_ok=True)
    safe_name   = project_name.replace(' ', '_').replace('/', '_')
    output_path = f"./outputs/SRS_{safe_name}_v{context['version']}.docx"
    doc.save(output_path)

    print(f"\nSRS generated successfully")
    print(f"   File     : {output_path}")
    print(f"   FRs      : {len(fr_list)}")
    nfr_cats = [k for k in grouped if k.startswith('NFR_') and grouped[k]]
    print(f"   NFR cats : {len(nfr_cats)}")
    if flagged:
        print(f"   Flagged  : {len(flagged)} (in Appendix A)")

    return output_path
